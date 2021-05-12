from __future__ import absolute_import
from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()

# Profile Picture
document.add_picture("profile.jpg" , width=Inches(2.0))

# name , phone_number and email details 
name = input("Whats your name : ")
speak("Hello" + name + "How are you today?")
speak("Whats your phone number")
phone_number = input("Whats your phone number : ")
speak("Whats your email")
email = input("Whats your email : ")

document.add_paragraph(
    name + ' | ' + phone_number + " | " + email
)

# about me 
document.add_heading("About me")
document.add_paragraph(input('Tell about yourself ? '))

# work experience
document.add_heading("Work Experience")
p = document.add_paragraph()

company = input('Enter company : ')
from_date = input('From Date : ')
to_date = input(' To_Date : ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input('describe your exeperience at ' + company)
p.add_run(experience_details)

# more experiences
while True :
    has_more_experiences = input("Do you have more experiences ? Yes Or No ")

    if has_more_experiences.lower() == 'yes' :
        p = document.add_paragraph()

        company = input('Enter company : ')
        from_date = input('From Date : ')
        to_date = input(' To_Date : ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        experience_details = input('describe your exeperience at ' + company)
        p.add_run(experience_details)

    else :
        break
# Skills 
document.add_heading("Skills")
document.add_paragraph(input("Enter you Skill"), style="List Bullet")

while True :
    has_more_skills = input("Do you have more skills ? Yes or No ")
    if has_more_skills.lower() == "yes" :
        document.add_paragraph(input("Enter you Skill"), style="List Bullet")
    else :
        break


# footer
section = document.sections[0]
footer = section.footer
footer_para = footer.paragraphs[0]
footer_para.text = "\tCV generated using Aamigoscode and Intuit QuickBooks course projects "
  


document.save('cv.docx')